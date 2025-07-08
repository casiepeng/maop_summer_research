# imports here
import requests
from docx import Document
import google.generativeai as genai
#---------------------------------------------------------------------------------------------------
# Prompts specifically chatGPT -4 Model API 
#
#
#
#
#
#
#
#---------------------------------------------------------------------------------------------------

# Prompt to ask all AIs
prompt = "Explain the significance of photosynthesis in under 150 words."

# ---- OpenAI GPT-4 ----
def query_openai(prompt):
    headers = {
        "Authorization": f"Bearer YOUR_OPENAI_API_KEY",
        "Content-Type": "application/json",
    }
    json_data = {
        "model": "gpt-4",
        "messages": [{"role": "user", "content": prompt}],
    }
    try:
        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=json_data)
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        return f"OpenAI error: {e}"

# ---- Claude 3 ----
def query_claude(prompt):
    headers = {
        "x-api-key": "YOUR_ANTHROPIC_API_KEY",
        "anthropic-version": "2023-06-01",
        "Content-Type": "application/json",
    }
    json_data = {
        "model": "claude-3-opus-20240229",
        "max_tokens": 1024,
        "messages": [{"role": "user", "content": prompt}],
    }
    try:
        response = requests.post("https://api.anthropic.com/v1/messages", headers=headers, json=json_data)
        return response.json()["content"][0]["text"]
    except Exception as e:
        return f"Claude error: {e}"

# ---- Gemini ----
def query_gemini(prompt):
    try:
        genai.configure(api_key="YOUR_GEMINI_API_KEY")
        model = genai.GenerativeModel("gemini-pro")
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Gemini error: {e}"

# ---- Save Responses to Word ----
def save_to_docx(responses, filename="ai_comparison.docx"):
    doc = Document()
    doc.add_heading("AI Responses Comparison", level=0)
    for model, response in responses.items():
        doc.add_heading(model, level=1)
        doc.add_paragraph(response)
    doc.save(filename)
    print(f"Saved to {filename}")

# ---- Main ----
if __name__ == "__main__":
    responses = {
        "OpenAI GPT-4": query_openai(prompt),
        "Anthropic Claude 3": query_claude(prompt),
        "Google Gemini": query_gemini(prompt),
    }

    save_to_docx(responses)